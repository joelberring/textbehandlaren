from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
import os
import io
import re
import requests
import socket
import ipaddress
from urllib.parse import urlparse
from datetime import datetime
from backend.app.services.template_parser import is_heading, is_instruction

class ExporterService:
    _IMAGE_SUGGESTION_RE = re.compile(r"\[BILDFÖRSLAG:\s*(.*?)\]", re.IGNORECASE)
    _ANSWER_INJECT_TOKEN = "__TB_INJECT_ANSWER__"
    _MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 8MB safeguard

    def __init__(self):
        self.output_dir = "exports"
        self.template_dir = "backend/app/templates"
        self.template_path = os.path.join(self.template_dir, "template.docx")
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)
        
        # Ensure a template exists
        if not os.path.exists(self.template_path):
            self._create_default_template()

    def _create_default_template(self):
        doc = Document()
        doc.add_heading('{{ title }}', 0)
        doc.add_paragraph('Datum: {{ date }}')
        
        doc.add_heading('Fråga', level=1)
        doc.add_paragraph('{{ query }}')
        
        doc.add_heading('Svar', level=1)
        doc.add_paragraph('{{ answer }}')
        
        doc.add_heading('Källhänvisningar', level=1)
        doc.add_paragraph('{% for source in sources %}')
        p = doc.add_paragraph('• Källa: {{ source.filename }} ({{ source.page }})', style='List Bullet')
        doc.add_paragraph('  {{ source.snippet }}')
        doc.add_paragraph('{% endfor %}')
        
        # V9: Images section placeholder
        doc.add_heading('Bilagor (Bilder)', level=1)
        doc.add_paragraph('{% for img in images %}')
        doc.add_paragraph('{{ img.inline_image }}')
        doc.add_paragraph('Beskrivning: {{ img.description }}')
        doc.add_paragraph('{% endfor %}')
        
        doc.save(self.template_path)

    def _template_has_placeholders(self, template_path: str, names: list) -> bool:
        try:
            doc = Document(template_path)
        except Exception:
            return False

        def has_name(text: str) -> bool:
            for name in names:
                if f"{{{{ {name} }}}}" in text or f"{{{{{name}}}}}" in text:
                    return True
            return False

        for p in doc.paragraphs:
            if has_name(p.text):
                return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if has_name(cell.text):
                        return True
        return False

    def _template_placeholder_names_present(self, template_path: str, names: list) -> set:
        try:
            doc = Document(template_path)
        except Exception:
            return set()

        found = set()

        def scan_text(text: str):
            for name in names:
                if f"{{{{ {name} }}}}" in text or f"{{{{{name}}}}}" in text:
                    found.add(name)

        for p in doc.paragraphs:
            scan_text(p.text or "")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        scan_text(p.text or "")

        return found

    def _extract_sections(self, markdown_text: str):
        sections = []
        current = None
        for line in markdown_text.splitlines():
            line = line.strip()
            if not line:
                if current:
                    current["content"] += "\n"
                continue
            if line.startswith("#"):
                # close previous
                if current:
                    sections.append(current)
                title = line.lstrip("#").strip()
                current = {"title": title, "content": ""}
            else:
                if current is None:
                    # preamble before first heading
                    current = {"title": "Inledning", "content": ""}
                current["content"] += line + "\n"
        if current:
            sections.append(current)
        return sections

    def _slugify(self, text: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9]+", "_", text.strip().lower())
        slug = re.sub(r"_+", "_", slug).strip("_")
        return slug or "section"

    def _pick_style_map(self, doc):
        available_styles = [s.name for s in doc.styles if s.type == 1]
        style_map = {
            'h1': next((s for s in ['Rubrik 1 med numrering', 'Rubrik 1', 'Rubrik 1 Huvudrubrik', 'Huvudrubrik', 'Heading 1', 'Heading 1 med numrering'] if s in available_styles), 'Heading 1'),
            'h2': next((s for s in ['Rubrik 2 med numrering', 'Rubrik 2', 'Underrubrik', 'Heading 2'] if s in available_styles), 'Heading 2'),
            'h3': next((s for s in ['Rubrik 3 med numrering', 'Rubrik 3', 'Heading 3'] if s in available_styles), 'Heading 3'),
            'normal': next((s for s in ['Svarstext', 'Brödtext', 'Normal', 'Body Text'] if s in available_styles), 'Normal'),
            'bullet': next((s for s in ['List Bullet', 'Punktlista', 'ListBullet'] if s in available_styles), 'List Bullet'),
            'caption': next((s for s in ['Bildtext', 'Caption'] if s in available_styles), None),
        }
        return style_map, available_styles

    def _add_inline_markdown_runs(self, paragraph: Paragraph, text: str):
        """
        Minimal inline markdown -> Word runs:
        - **bold**
        - *italic*
        Best-effort, non-nested.
        """
        if text is None:
            return
        s = str(text)
        i = 0
        bold = False
        italic = False

        while i < len(s):
            if s.startswith("**", i):
                bold = not bold
                i += 2
                continue
            if s[i] == "*":
                italic = not italic
                i += 1
                continue

            j = i
            while j < len(s) and (not s.startswith("**", j)) and s[j] != "*":
                j += 1
            seg = s[i:j]
            if seg:
                run = paragraph.add_run(seg)
                run.bold = bool(bold)
                run.italic = bool(italic)
            i = j

    def _iter_all_paragraphs(self, doc):
        # Body paragraphs
        for p in doc.paragraphs:
            yield p
        # Table paragraphs (including nested tables)
        def walk_table(table):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for t in cell.tables:
                        yield from walk_table(t)
        for table in doc.tables:
            yield from walk_table(table)

    def _inject_markdown_after(self, anchor: Paragraph, markdown_text: str, style_map: dict, available_styles: list) -> int:
        last = anchor
        inserted = 0
        for raw_line in (markdown_text or "").splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()
            if not stripped:
                continue

            style = style_map.get('normal') or 'Normal'
            content = stripped

            # Headings
            if stripped.startswith('### '):
                style = style_map.get('h3') or style
                content = stripped[4:].strip()
            elif stripped.startswith('## '):
                style = style_map.get('h2') or style
                content = stripped[3:].strip()
            elif stripped.startswith('# '):
                style = style_map.get('h1') or style
                content = stripped[2:].strip()
            # Bullets
            elif stripped.startswith('- ') or stripped.startswith('* '):
                style = style_map.get('bullet') or style
                content = stripped[2:].strip()
            # Numbered lists
            elif re.match(r'^\d+\.', stripped):
                content = re.sub(r'^\d+\.\s*', '', stripped).strip()
                style = 'List Number' if 'List Number' in available_styles else style
            # Captions (best-effort)
            elif style_map.get('caption') and re.match(r'^(bildtext|figur|tabell)\s*:', stripped, re.IGNORECASE):
                style = style_map.get('caption') or style

            p = self._insert_paragraph_after(last, "", style=style)
            self._add_inline_markdown_runs(p, content)
            last = p
            inserted += 1

        return inserted

    def _inject_markdown_at_token(self, doc, token: str, markdown_text: str) -> int:
        if not token:
            return 0
        style_map, available_styles = self._pick_style_map(doc)

        for p in self._iter_all_paragraphs(doc):
            if token not in (p.text or ""):
                continue

            inserted = self._inject_markdown_after(p, markdown_text, style_map, available_styles)

            # Remove/clean anchor paragraph
            try:
                if (p.text or "").strip() == token:
                    self._remove_paragraph(p)
                else:
                    p.text = (p.text or "").replace(token, "").strip()
            except Exception:
                pass
            return inserted
        return 0

    def _append_source_appendix(self, doc, sources: list):
        if not sources:
            return

        available_styles = [s.name for s in doc.styles if s.type == 1]
        h1_style = next((s for s in ['Rubrik 1 med numrering', 'Rubrik 1', 'Huvudrubrik', 'Heading 1'] if s in available_styles), 'Heading 1')
        h2_style = next((s for s in ['Rubrik 2 med numrering', 'Rubrik 2', 'Underrubrik', 'Heading 2'] if s in available_styles), 'Heading 2')
        normal_style = next((s for s in ['Svarstext', 'Brödtext', 'Normal', 'Body Text'] if s in available_styles), 'Normal')

        unique = []
        seen = set()
        for s in sources:
            meta = s.get("metadata", {})
            key = (
                s.get("source_ref") or meta.get("source_ref") or "",
                meta.get("filename", ""),
                str(meta.get("page", "")),
                meta.get("library_id", "")
            )
            if key in seen:
                continue
            seen.add(key)
            unique.append(s)

        if not unique:
            return

        doc.add_paragraph("")
        doc.add_paragraph("Källbilaga (spårbarhet)", style=h1_style)
        doc.add_paragraph(
            "Nedan listas använda källor med käll-ID för spårbarhet till AI-svaret.",
            style=normal_style
        )

        for s in unique:
            meta = s.get("metadata", {})
            source_ref = s.get("source_ref") or meta.get("source_ref") or "-"
            filename = meta.get("filename", "Okänd fil")
            page = meta.get("page", "-")
            library_name = meta.get("library_name", "Okänt bibliotek")
            library_type = meta.get("library_type", s.get("type", "OKÄND"))
            doc_id = meta.get("doc_id", "-")
            snippet = (s.get("content") or "").strip()
            if len(snippet) > 450:
                snippet = snippet[:450] + "..."

            doc.add_paragraph(f"Källa {source_ref}", style=h2_style)
            doc.add_paragraph(f"Fil: {filename}", style=normal_style)
            doc.add_paragraph(f"Sida: {page}", style=normal_style)
            doc.add_paragraph(f"Bibliotek: {library_name} ({library_type})", style=normal_style)
            doc.add_paragraph(f"Dokument-ID: {doc_id}", style=normal_style)
            doc.add_paragraph(f"Utdrag: {snippet}", style=normal_style)

    def _tokenize(self, text: str):
        return re.findall(r"[a-z0-9åäö\-]{3,}", (text or "").lower())

    def _parse_image_suggestion(self, suggestion_text: str):
        parts = [p.strip() for p in (suggestion_text or "").split("|")]
        parts += ["", "", "", ""]
        what = parts[0]
        source = parts[1]
        page_text = parts[2]
        section = parts[3]

        page_num = None
        match = re.search(r"(\d{1,4})", page_text)
        if match:
            try:
                page_num = int(match.group(1))
            except Exception:
                page_num = None

        return {
            "raw": suggestion_text.strip(),
            "what": what.strip(),
            "source": source.strip().lower(),
            "section": section.strip(),
            "page": page_num,
        }

    def _score_image_candidate(self, suggestion: dict, image: dict, already_used_ids: set):
        source_document = str(image.get("source_document", ""))
        page = image.get("page")
        tags = image.get("tags") or []
        section_hints = image.get("section_hints") or []
        desc = image.get("description", "")
        context = image.get("context_excerpt", "")

        haystack = " ".join([
            source_document,
            str(page or ""),
            desc,
            context,
            " ".join([str(t) for t in tags]),
            " ".join([str(h) for h in section_hints]),
        ]).lower()

        suggestion_tokens = self._tokenize(
            f"{suggestion.get('what', '')} {suggestion.get('section', '')}"
        )
        overlap = sum(1 for token in suggestion_tokens if token in haystack)
        score = overlap * 3

        source_hint = suggestion.get("source")
        if source_hint:
            if source_hint in source_document.lower():
                score += 8
            else:
                score -= 1

        page_hint = suggestion.get("page")
        if page_hint is not None:
            try:
                if int(page or -1) == int(page_hint):
                    score += 8
                else:
                    score -= 1
            except Exception:
                pass

        image_id = image.get("id") or f"{source_document}:{page}:{image.get('url', '')}"
        if image_id in already_used_ids:
            score -= 4

        try:
            score += int(image.get("library_priority", 50)) // 30
        except Exception:
            pass

        return score

    def _select_image_for_suggestion(self, suggestion_text: str, matched_images: list, already_used_ids: set):
        if not matched_images:
            return None
        suggestion = self._parse_image_suggestion(suggestion_text)

        best = None
        best_score = -10**9
        for img in matched_images:
            score = self._score_image_candidate(suggestion, img, already_used_ids)
            if score > best_score:
                best_score = score
                best = img

        if best is None:
            return None

        # Require minimal relevance when a concrete hint exists.
        has_specific_hint = bool(suggestion.get("source") or suggestion.get("page") is not None or suggestion.get("what"))
        if has_specific_hint and best_score < 1:
            return None
        return best

    def _is_safe_remote_url(self, url: str) -> bool:
        """
        Basic SSRF protection for server-side fetching.
        We only allow http(s) and block local/private/link-local/loopback targets.
        """
        if not url:
            return False
        try:
            parsed = urlparse(url)
        except Exception:
            return False

        if parsed.scheme not in ("https", "http"):
            return False
        if not parsed.netloc:
            return False
        if parsed.username or parsed.password:
            return False

        host = (parsed.hostname or "").strip().lower()
        if not host:
            return False
        if host in {"localhost"} or host.endswith(".local"):
            return False

        try:
            # Resolve and block any private-ish targets (including metadata IPs).
            infos = socket.getaddrinfo(host, None)
            for info in infos:
                ip_str = info[4][0]
                ip = ipaddress.ip_address(ip_str)
                if (
                    ip.is_private
                    or ip.is_loopback
                    or ip.is_link_local
                    or ip.is_multicast
                    or ip.is_reserved
                    or ip.is_unspecified
                ):
                    return False
        except Exception:
            return False

        return True

    def _download_image_bytes(self, url: str):
        if not url:
            return None
        if not self._is_safe_remote_url(url):
            print(f"Blocked unsafe image URL: {url}")
            return None
        try:
            response = requests.get(url, timeout=12, stream=True)
            if response.status_code != 200:
                return None

            content_length = response.headers.get("Content-Length")
            if content_length:
                try:
                    if int(content_length) > self._MAX_IMAGE_BYTES:
                        print(f"Blocked large image download ({content_length} bytes): {url}")
                        return None
                except Exception:
                    pass

            buf = bytearray()
            for chunk in response.iter_content(chunk_size=64 * 1024):
                if not chunk:
                    continue
                buf.extend(chunk)
                if len(buf) > self._MAX_IMAGE_BYTES:
                    print(f"Blocked large image download (> {self._MAX_IMAGE_BYTES} bytes): {url}")
                    return None
            return bytes(buf) if buf else None
        except Exception as e:
            print(f"Failed to download image {url}: {e}")
        return None

    def _replace_image_suggestions(self, doc, matched_images: list):
        if not matched_images:
            return 0

        available_styles = [s.name for s in doc.styles if s.type == 1]
        normal_style = next((s for s in ['Svarstext', 'Brödtext', 'Normal', 'Body Text'] if s in available_styles), 'Normal')
        caption_style = next((s for s in ['Bildtext', 'Caption'] if s in available_styles), normal_style)

        inserted = 0
        used_ids = set()
        paragraphs = list(self._iter_all_paragraphs(doc))

        for para in paragraphs:
            text = para.text or ""
            matches = list(self._IMAGE_SUGGESTION_RE.finditer(text))
            if not matches:
                continue

            fallback_notes = []
            anchor = para
            for match in matches:
                suggestion_raw = (match.group(1) or "").strip()
                if not suggestion_raw:
                    continue
                image = self._select_image_for_suggestion(suggestion_raw, matched_images, used_ids)
                if not image:
                    fallback_notes.append(f"Bildförslag: {suggestion_raw}")
                    continue

                image_bytes = self._download_image_bytes(image.get("url"))
                if not image_bytes:
                    fallback_notes.append(f"Bildförslag: {suggestion_raw}")
                    continue

                image_para = self._insert_paragraph_after(anchor)
                image_para.add_run().add_picture(io.BytesIO(image_bytes), width=Mm(130))

                caption = (
                    f"Figur: {image.get('description', 'Bild utan beskrivning')[:180]} "
                    f"(Källa: {image.get('source_document', 'okänd')}, sida {image.get('page', '-')})"
                )
                anchor = self._insert_paragraph_after(image_para, caption, style=caption_style)

                image_id = image.get("id") or f"{image.get('source_document')}:{image.get('page')}:{image.get('url', '')}"
                used_ids.add(image_id)
                inserted += 1

            cleaned = self._IMAGE_SUGGESTION_RE.sub("", text).strip()
            if fallback_notes:
                if cleaned:
                    cleaned = f"{cleaned}\n" + "\n".join(fallback_notes)
                else:
                    cleaned = "\n".join(fallback_notes)

            if cleaned:
                para.text = cleaned
            else:
                self._remove_paragraph(para)

        return inserted

    async def generate_word_response(self, query: str, answer: str, sources: list, 
                                      template_path: str = None, matched_images: list = None):
        path_to_use = template_path or self.template_path
        if not os.path.exists(path_to_use):
            self._create_default_template()
            path_to_use = self.template_path

        doc = DocxTemplate(path_to_use)

        # If the template uses a plain-text placeholder for the answer, inject structured markdown at that location
        # instead of rendering the whole answer into a single paragraph.
        answer_placeholders = self._template_placeholder_names_present(
            path_to_use,
            ["answer", "content", "ai_text"]
        )
        inject_token = self._ANSWER_INJECT_TOKEN if answer_placeholders else None
        
        # Prepare context for template
        formatted_sources = []
        for s in sources:
            metadata = s.get('metadata', {})
            formatted_sources.append({
                'source_ref': s.get('source_ref', metadata.get('source_ref')),
                'filename': metadata.get('filename', 'Okänd'),
                'page': metadata.get('page', 'N/A'),
                'library_name': metadata.get('library_name', 'Okänt bibliotek'),
                'library_type': metadata.get('library_type', s.get('type', 'OKÄND')),
                'doc_id': metadata.get('doc_id', '-'),
                'snippet': s.get('content', '')[:300] + "..."
            })

        # V9: Process matched images
        has_image_suggestions = bool(self._IMAGE_SUGGESTION_RE.search(answer or ""))
        formatted_images = []
        if matched_images and not has_image_suggestions:
            for img in matched_images[:5]:  # Limit to 5 images
                try:
                    # Download image from URL (SSRF-protected)
                    image_bytes = self._download_image_bytes(img.get('url'))
                    if image_bytes:
                        image_stream = io.BytesIO(image_bytes)
                        inline_image = InlineImage(doc, image_stream, width=Mm(100))
                        formatted_images.append({
                            'inline_image': inline_image,
                            'description': img.get('description', 'Ingen beskrivning')[:200],
                            'source': img.get('source_document', 'Okänd källa')
                        })
                except Exception as e:
                    print(f"Failed to download image: {e}")
                    continue

        sections = self._extract_sections(answer)
        section_map = {}
        for sec in sections:
            key = f"section_{self._slugify(sec['title'])}"
            section_map[key] = sec["content"].strip()

        context = {
            'title': 'Textbehandlaren Export',
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'query': query,
            'answer': inject_token or answer, # Fallback for old templates
            'content': inject_token or answer,
            'ai_text': inject_token or answer,
            'sources': formatted_sources,
            'images': formatted_images,
            'sections': sections,
            **section_map
        }
        
        doc.render(context)
        
        injected = 0
        if inject_token:
            try:
                injected = self._inject_markdown_at_token(doc, inject_token, answer)
            except Exception as e:
                print(f"Placeholder injection failed: {e}")

        if not inject_token:
            # If the template contains headings with instruction text, replace per section.
            try:
                injected = self._apply_sections_by_headings(doc, answer) or 0
            except Exception as e:
                print(f"Section injection failed: {e}")

        # Fallback: append a styled version of the answer at the end.
        if injected == 0:
            self._inject_styled_answer(doc, answer)

        # Replace explicit [BILDFÖRSLAG: ...] markers with actual images in-place.
        try:
            inserted = self._replace_image_suggestions(doc, matched_images or [])
            if inserted:
                print(f"Inserted {inserted} image(s) from BILDFÖRSLAG markers.")
        except Exception as e:
            print(f"Image suggestion replacement failed: {e}")

        # Always append a traceable source appendix for auditability.
        try:
            self._append_source_appendix(doc, sources)
        except Exception as e:
            print(f"Source appendix generation failed: {e}")
        
        filename = f"Textbehandlaren_Export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        return file_path

    def _inject_styled_answer(self, doc, markdown_text):
        """Append or inject the AI answer with appropriate template styles."""
        style_map, available_styles = self._pick_style_map(doc)

        # Simple Markdown line-by-line parser
        lines = markdown_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headings
            if line.startswith('### '):
                p = doc.add_paragraph("", style=style_map['h3'])
                self._add_inline_markdown_runs(p, line[4:].strip())
            elif line.startswith('## '):
                p = doc.add_paragraph("", style=style_map['h2'])
                self._add_inline_markdown_runs(p, line[3:].strip())
            elif line.startswith('# '):
                p = doc.add_paragraph("", style=style_map['h1'])
                self._add_inline_markdown_runs(p, line[2:].strip())
            # Bullets
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph("", style=style_map['bullet'])
                self._add_inline_markdown_runs(p, line[2:].strip())
            # Numbered lists (simplified)
            elif re.match(r'^\d+\.', line):
                content = re.sub(r'^\d+\.\s*', '', line)
                p = doc.add_paragraph("", style='List Number' if 'List Number' in available_styles else style_map['normal'])
                self._add_inline_markdown_runs(p, content.strip())
            # Captions (best-effort)
            elif style_map.get('caption') and re.match(r'^(bildtext|figur|tabell)\s*:', line, re.IGNORECASE):
                p = doc.add_paragraph("", style=style_map['caption'])
                self._add_inline_markdown_runs(p, line)
            # Normal text
            else:
                p = doc.add_paragraph("", style=style_map['normal'])
                self._add_inline_markdown_runs(p, line)

    def _remove_paragraph(self, paragraph: Paragraph):
        p = paragraph._element
        p.getparent().remove(p)

    def _insert_paragraph_after(self, paragraph: Paragraph, text: str = "", style: str = None) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if style:
            new_para.style = style
        if text:
            new_para.add_run(text)
        return new_para

    def _normalize_heading_key(self, text: str) -> str:
        s = (text or "").strip().lower()
        s = re.sub(r"\s+", " ", s).strip()
        # Remove leading numbering like "1.", "1.2", "1)" etc.
        s = re.sub(r"^\d+(?:\.\d+)*\s*[\)\.\-:]*\s*", "", s).strip()
        s = s.rstrip(":").strip()
        return s

    def _apply_sections_by_headings(self, doc, answer_markdown: str):
        style_map, available_styles = self._pick_style_map(doc)

        sections = self._extract_sections(answer_markdown)
        section_by_title = {self._normalize_heading_key(s["title"]): s["content"].strip() for s in sections}

        paragraphs = list(self._iter_all_paragraphs(doc))
        inserted_total = 0
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            if is_heading(para):
                heading_title = self._normalize_heading_key((para.text or "").strip())
                # Remove instruction paragraphs directly under heading
                j = i + 1
                while j < len(paragraphs) and not is_heading(paragraphs[j]):
                    if is_instruction(paragraphs[j]):
                        self._remove_paragraph(paragraphs[j])
                    j += 1

                # Insert generated content after heading
                content = section_by_title.get(heading_title)
                if content:
                    last = para
                    for line in content.splitlines():
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith("### "):
                            last = self._insert_paragraph_after(last, "", style_map['h3'])
                            self._add_inline_markdown_runs(last, line[4:].strip())
                        elif line.startswith("## "):
                            last = self._insert_paragraph_after(last, "", style_map['h2'])
                            self._add_inline_markdown_runs(last, line[3:].strip())
                        elif line.startswith("# "):
                            last = self._insert_paragraph_after(last, "", style_map['h1'])
                            self._add_inline_markdown_runs(last, line[2:].strip())
                        elif line.startswith("- ") or line.startswith("* "):
                            last = self._insert_paragraph_after(last, "", style_map['bullet'])
                            self._add_inline_markdown_runs(last, line[2:].strip())
                        elif re.match(r'^\d+\.', line):
                            content_line = re.sub(r'^\d+\.\s*', '', line)
                            last = self._insert_paragraph_after(
                                last, "",
                                'List Number' if 'List Number' in available_styles else None
                            )
                            self._add_inline_markdown_runs(last, content_line.strip())
                        else:
                            last = self._insert_paragraph_after(last, "", style_map['normal'])
                            self._add_inline_markdown_runs(last, line)
                        inserted_total += 1
                i = j
            else:
                i += 1
        return inserted_total

exporter_service = ExporterService()
